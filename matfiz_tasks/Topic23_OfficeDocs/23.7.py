import docx
import re

pattern2 = r"\d{2}\.\d{2}\.\d{4}"

doc=docx.Document('123.docx')
l=list()

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        t=run.text
        #print(t)
        if (re.findall(pattern2, t) != []):
            l.append(re.findall(pattern2, t))
            print('в цьому параграфі є текс', t)
print(l)
doc=docx.Document()
for i in l:
    doc.add_paragraph(text=i,style='Body Text')
fname='1234.docx'
doc.save(fname)
